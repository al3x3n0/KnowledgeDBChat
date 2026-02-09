#!/usr/bin/env python3
"""
Script to create sample DOCX templates for the Template Fill feature.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(OUTPUT_DIR, "templates")


def create_project_report_template():
    """Create a Project Status Report template."""
    doc = Document()

    # Title
    title = doc.add_heading("Project Status Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("[Project Name]")
    run.italic = True

    doc.add_paragraph()  # Spacer

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("[Provide a brief overview of the project status, key achievements, and critical issues.]")

    # Project Overview
    doc.add_heading("Project Overview", level=1)
    doc.add_paragraph("[Describe the project goals, scope, and current phase.]")

    # Technical Progress
    doc.add_heading("Technical Progress", level=1)

    doc.add_heading("Completed Milestones", level=2)
    doc.add_paragraph("[List recently completed technical milestones and achievements.]")

    doc.add_heading("Current Development Status", level=2)
    doc.add_paragraph("[Describe the current state of development, including features being worked on.]")

    doc.add_heading("Technical Challenges", level=2)
    doc.add_paragraph("[Document any technical challenges encountered and their resolutions.]")

    # Architecture and Design
    doc.add_heading("Architecture and Design", level=1)
    doc.add_paragraph("[Provide an overview of the system architecture and key design decisions.]")

    # Performance Metrics
    doc.add_heading("Performance Metrics", level=1)
    doc.add_paragraph("[Include relevant performance benchmarks and metrics.]")

    # Risk Assessment
    doc.add_heading("Risk Assessment", level=1)
    doc.add_paragraph("[Identify current project risks and mitigation strategies.]")

    # Next Steps
    doc.add_heading("Next Steps", level=1)
    doc.add_paragraph("[Outline planned activities for the next reporting period.]")

    # Save
    filepath = os.path.join(TEMPLATES_DIR, "project_report_template.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_technical_specification_template():
    """Create a Technical Specification template."""
    doc = Document()

    # Title
    title = doc.add_heading("Technical Specification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Document Info Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    cells = [
        ("Document Version:", "[1.0]"),
        ("Author:", "[Author Name]"),
        ("Date:", "[Date]"),
        ("Status:", "[Draft/Review/Approved]"),
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Introduction
    doc.add_heading("Introduction", level=1)

    doc.add_heading("Purpose", level=2)
    doc.add_paragraph("[Describe the purpose of this technical specification.]")

    doc.add_heading("Scope", level=2)
    doc.add_paragraph("[Define the scope and boundaries of the system being specified.]")

    doc.add_heading("Definitions and Acronyms", level=2)
    doc.add_paragraph("[List key terms and acronyms used in this document.]")

    # System Overview
    doc.add_heading("System Overview", level=1)
    doc.add_paragraph("[Provide a high-level description of the system.]")

    # Architecture
    doc.add_heading("System Architecture", level=1)

    doc.add_heading("Architecture Overview", level=2)
    doc.add_paragraph("[Describe the overall system architecture.]")

    doc.add_heading("Component Diagram", level=2)
    doc.add_paragraph("[Include or describe the component architecture.]")

    doc.add_heading("Data Flow", level=2)
    doc.add_paragraph("[Describe how data flows through the system.]")

    # Components
    doc.add_heading("Component Specifications", level=1)

    doc.add_heading("Core Components", level=2)
    doc.add_paragraph("[Detail the core components and their responsibilities.]")

    doc.add_heading("External Interfaces", level=2)
    doc.add_paragraph("[Describe external interfaces and integrations.]")

    # API Specification
    doc.add_heading("API Specification", level=1)
    doc.add_paragraph("[Document the API endpoints, parameters, and responses.]")

    # Performance Requirements
    doc.add_heading("Performance Requirements", level=1)
    doc.add_paragraph("[Specify performance requirements and benchmarks.]")

    # Security Considerations
    doc.add_heading("Security Considerations", level=1)
    doc.add_paragraph("[Document security requirements and measures.]")

    # Dependencies
    doc.add_heading("Dependencies", level=1)
    doc.add_paragraph("[List external dependencies and version requirements.]")

    # Save
    filepath = os.path.join(TEMPLATES_DIR, "technical_specification_template.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_executive_summary_template():
    """Create an Executive Summary template."""
    doc = Document()

    # Title
    title = doc.add_heading("Executive Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project name
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("[Project/Product Name]")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Overview
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("[Provide a concise overview of the project/product and its purpose.]")

    # Key Features
    doc.add_heading("Key Features and Capabilities", level=1)
    doc.add_paragraph("[Highlight the main features and capabilities.]")

    # Business Value
    doc.add_heading("Business Value", level=1)
    doc.add_paragraph("[Describe the business value and benefits.]")

    # Technical Highlights
    doc.add_heading("Technical Highlights", level=1)
    doc.add_paragraph("[Summarize the key technical achievements and innovations.]")

    # Current Status
    doc.add_heading("Current Status", level=1)
    doc.add_paragraph("[Describe the current development/deployment status.]")

    # Key Metrics
    doc.add_heading("Key Metrics", level=1)
    doc.add_paragraph("[Include relevant performance or success metrics.]")

    # Roadmap
    doc.add_heading("Roadmap", level=1)
    doc.add_paragraph("[Outline the future development roadmap and planned features.]")

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("[Provide a concluding summary and recommendations.]")

    # Save
    filepath = os.path.join(TEMPLATES_DIR, "executive_summary_template.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_release_notes_template():
    """Create a Release Notes template."""
    doc = Document()

    # Title
    title = doc.add_heading("Release Notes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Version
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version [X.Y.Z]")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Release Overview
    doc.add_heading("Release Overview", level=1)
    doc.add_paragraph("[Provide a summary of this release and its significance.]")

    # New Features
    doc.add_heading("New Features", level=1)
    doc.add_paragraph("[List and describe new features introduced in this release.]")

    # Improvements
    doc.add_heading("Improvements", level=1)
    doc.add_paragraph("[Document improvements and enhancements to existing functionality.]")

    # Bug Fixes
    doc.add_heading("Bug Fixes", level=1)
    doc.add_paragraph("[List bugs that have been fixed in this release.]")

    # Performance Improvements
    doc.add_heading("Performance Improvements", level=1)
    doc.add_paragraph("[Describe any performance optimizations included.]")

    # Breaking Changes
    doc.add_heading("Breaking Changes", level=1)
    doc.add_paragraph("[Document any breaking changes and migration steps.]")

    # Known Issues
    doc.add_heading("Known Issues", level=1)
    doc.add_paragraph("[List any known issues in this release.]")

    # Upgrade Instructions
    doc.add_heading("Upgrade Instructions", level=1)
    doc.add_paragraph("[Provide instructions for upgrading from previous versions.]")

    # Save
    filepath = os.path.join(TEMPLATES_DIR, "release_notes_template.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


if __name__ == "__main__":
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    print("Creating DOCX templates...")
    create_project_report_template()
    create_technical_specification_template()
    create_executive_summary_template()
    create_release_notes_template()
    print("\nAll templates created successfully!")
