"""
DOCX document builder service.

Creates DOCX files from structured content using python-docx.
"""

from io import BytesIO
from typing import Dict, Optional, Any, List
from loguru import logger

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color string to RGBColor."""
    hex_color = hex_color.lstrip('#')
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


class DOCXBuilder:
    """
    Builds DOCX documents from structured content.

    Supports multiple built-in styles, custom themes, and various content types.
    """

    # Built-in style definitions
    STYLES = {
        "professional": {
            "title_color": "#1a365d",
            "heading_color": "#2e86ab",
            "text_color": "#333333",
            "accent_color": "#2e86ab",
            "title_font": "Calibri",
            "body_font": "Calibri",
            "code_font": "Consolas",
            "title_size": 28,
            "heading1_size": 18,
            "heading2_size": 14,
            "body_size": 11,
            "code_size": 10,
        },
        "casual": {
            "title_color": "#4a90d9",
            "heading_color": "#ff6b6b",
            "text_color": "#2d3a4a",
            "accent_color": "#ff6b6b",
            "title_font": "Arial",
            "body_font": "Arial",
            "code_font": "Courier New",
            "title_size": 30,
            "heading1_size": 20,
            "heading2_size": 16,
            "body_size": 12,
            "code_size": 10,
        },
        "technical": {
            "title_color": "#007acc",
            "heading_color": "#28a745",
            "text_color": "#24292e",
            "accent_color": "#28a745",
            "title_font": "Segoe UI",
            "body_font": "Segoe UI",
            "code_font": "Consolas",
            "title_size": 26,
            "heading1_size": 16,
            "heading2_size": 13,
            "body_size": 11,
            "code_size": 9,
        },
    }

    def __init__(
        self,
        style: str = "professional",
        custom_theme: Optional[Dict[str, Any]] = None
    ):
        """
        Initialize the builder with a style or custom theme.

        Args:
            style: Built-in style name (used if no custom_theme)
            custom_theme: Custom theme configuration dict
        """
        self.style = style

        if custom_theme:
            self.style_config = self._parse_custom_theme(custom_theme)
        else:
            self.style_config = self.STYLES.get(style, self.STYLES["professional"])

    def _parse_custom_theme(self, theme: Dict[str, Any]) -> Dict[str, Any]:
        """
        Parse custom theme configuration into style config.

        Args:
            theme: Custom theme dict

        Returns:
            Style config dict compatible with built-in styles
        """
        # Start with professional defaults
        config = dict(self.STYLES["professional"])

        # Override with custom values
        for key in theme:
            if key in config:
                config[key] = theme[key]

        return config

    @classmethod
    def get_available_styles(cls) -> Dict[str, Dict[str, str]]:
        """
        Get list of available built-in styles with descriptions.

        Returns:
            Dict mapping style name to description
        """
        return {
            "professional": {
                "name": "Professional",
                "description": "Clean, corporate look with dark blue accents",
                "primary_color": "#1a365d",
            },
            "casual": {
                "name": "Casual",
                "description": "Friendly and approachable with warm colors",
                "primary_color": "#4a90d9",
            },
            "technical": {
                "name": "Technical",
                "description": "Developer-focused with monospace code blocks",
                "primary_color": "#007acc",
            },
        }

    def build(
        self,
        title: str,
        content_items: List[Dict[str, Any]],
        author: Optional[str] = None,
        subject: Optional[str] = None
    ) -> bytes:
        """
        Build a DOCX document from content items.

        Args:
            title: Document title
            content_items: List of content items with type and data
            author: Optional author name for metadata
            subject: Optional subject for metadata

        Returns:
            DOCX file as bytes

        Content item types:
            - heading: {"type": "heading", "level": 1|2|3, "text": "..."}
            - paragraph: {"type": "paragraph", "text": "..."}
            - bullet_list: {"type": "bullet_list", "items": ["...", "..."]}
            - numbered_list: {"type": "numbered_list", "items": ["...", "..."]}
            - code_block: {"type": "code_block", "code": "...", "language": "python"}
            - table: {"type": "table", "headers": [...], "rows": [[...], ...]}
            - quote: {"type": "quote", "text": "..."}
            - horizontal_rule: {"type": "horizontal_rule"}
            - page_break: {"type": "page_break"}
        """
        doc = Document()

        # Set document properties
        core_props = doc.core_properties
        core_props.title = title
        if author:
            core_props.author = author
        if subject:
            core_props.subject = subject

        # Configure styles
        self._configure_styles(doc)

        # Add title
        self._add_title(doc, title)

        # Process content items
        for item in content_items:
            item_type = item.get("type", "paragraph")

            if item_type == "heading":
                self._add_heading(doc, item.get("text", ""), item.get("level", 1))
            elif item_type == "paragraph":
                self._add_paragraph(doc, item.get("text", ""))
            elif item_type == "bullet_list":
                self._add_bullet_list(doc, item.get("items", []))
            elif item_type == "numbered_list":
                self._add_numbered_list(doc, item.get("items", []))
            elif item_type == "code_block":
                self._add_code_block(doc, item.get("code", ""), item.get("language", ""))
            elif item_type == "table":
                self._add_table(doc, item.get("headers", []), item.get("rows", []))
            elif item_type == "quote":
                self._add_quote(doc, item.get("text", ""))
            elif item_type == "horizontal_rule":
                self._add_horizontal_rule(doc)
            elif item_type == "page_break":
                doc.add_page_break()
            else:
                # Default to paragraph
                self._add_paragraph(doc, str(item.get("text", item)))

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _configure_styles(self, doc: Document):
        """Configure document styles based on theme."""
        styles = doc.styles

        # Configure Normal style
        normal = styles['Normal']
        normal.font.name = self.style_config["body_font"]
        normal.font.size = Pt(self.style_config["body_size"])
        normal.font.color.rgb = hex_to_rgb(self.style_config["text_color"])

        # Configure Heading 1
        h1 = styles['Heading 1']
        h1.font.name = self.style_config["title_font"]
        h1.font.size = Pt(self.style_config["heading1_size"])
        h1.font.bold = True
        h1.font.color.rgb = hex_to_rgb(self.style_config["heading_color"])

        # Configure Heading 2
        h2 = styles['Heading 2']
        h2.font.name = self.style_config["title_font"]
        h2.font.size = Pt(self.style_config["heading2_size"])
        h2.font.bold = True
        h2.font.color.rgb = hex_to_rgb(self.style_config["heading_color"])

        # Try to configure Heading 3
        try:
            h3 = styles['Heading 3']
            h3.font.name = self.style_config["title_font"]
            h3.font.size = Pt(self.style_config["body_size"] + 1)
            h3.font.bold = True
            h3.font.color.rgb = hex_to_rgb(self.style_config["heading_color"])
        except KeyError:
            pass

    def _add_title(self, doc: Document, title: str):
        """Add document title."""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = self.style_config["title_font"]
        run.font.size = Pt(self.style_config["title_size"])
        run.font.bold = True
        run.font.color.rgb = hex_to_rgb(self.style_config["title_color"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(24)

        # Add a line under title
        self._add_horizontal_rule(doc)
        doc.add_paragraph()  # Space after rule

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add a heading."""
        level = min(max(level, 1), 3)  # Clamp to 1-3
        doc.add_heading(text, level=level)

    def _add_paragraph(self, doc: Document, text: str):
        """Add a paragraph."""
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(12)

    def _add_bullet_list(self, doc: Document, items: List[str]):
        """Add a bullet list."""
        for item in items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(6)

    def _add_numbered_list(self, doc: Document, items: List[str]):
        """Add a numbered list."""
        for item in items:
            p = doc.add_paragraph(item, style='List Number')
            p.paragraph_format.space_after = Pt(6)

    def _add_code_block(self, doc: Document, code: str, language: str = ""):
        """Add a code block with monospace formatting."""
        p = doc.add_paragraph()

        # Add language label if provided
        if language:
            label_run = p.add_run(f"{language}\n")
            label_run.font.name = self.style_config["code_font"]
            label_run.font.size = Pt(self.style_config["code_size"] - 1)
            label_run.font.bold = True
            label_run.font.color.rgb = hex_to_rgb(self.style_config["accent_color"])

        # Add code content
        code_run = p.add_run(code)
        code_run.font.name = self.style_config["code_font"]
        code_run.font.size = Pt(self.style_config["code_size"])
        code_run.font.color.rgb = hex_to_rgb(self.style_config["text_color"])

        # Add background shading
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F5F5F5')  # Light gray background
        pPr.append(shd)

        # Set paragraph formatting
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    def _add_table(self, doc: Document, headers: List[str], rows: List[List[str]]):
        """Add a table."""
        if not headers and not rows:
            return

        num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
        if num_cols == 0:
            return

        table = doc.add_table(rows=1 if headers else 0, cols=num_cols)
        table.style = 'Table Grid'

        # Add header row
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                # Style header cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = hex_to_rgb(self.style_config["heading_color"])
                # Add background shading to header
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E8E8E8')
                tcPr.append(shd)

        # Add data rows
        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                if i < num_cols:
                    row.cells[i].text = str(cell_text)

        # Add space after table
        doc.add_paragraph()

    def _add_quote(self, doc: Document, text: str):
        """Add a blockquote."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = hex_to_rgb(self.style_config["text_color"])

        # Add left border effect via indentation
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

        # Add vertical bar effect
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')  # border width
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), self.style_config["accent_color"].lstrip('#'))
        pBdr.append(left)
        pPr.append(pBdr)

    def _add_horizontal_rule(self, doc: Document):
        """Add a horizontal rule."""
        p = doc.add_paragraph()
        p_element = p._element
        pPr = p_element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), self.style_config["accent_color"].lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)


def markdown_to_content_items(markdown_text: str) -> List[Dict[str, Any]]:
    """
    Convert markdown text to content items for DOCXBuilder.

    Args:
        markdown_text: Markdown formatted text

    Returns:
        List of content items
    """
    import re

    content_items = []
    lines = markdown_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            content_items.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # Code block
        if line.strip().startswith('```'):
            language = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            content_items.append({
                "type": "code_block",
                "code": '\n'.join(code_lines),
                "language": language
            })
            i += 1
            continue

        # Bullet list item
        bullet_match = re.match(r'^[\*\-\+]\s+(.+)$', line)
        if bullet_match:
            items = [bullet_match.group(1)]
            i += 1
            while i < len(lines):
                next_bullet = re.match(r'^[\*\-\+]\s+(.+)$', lines[i])
                if next_bullet:
                    items.append(next_bullet.group(1))
                    i += 1
                else:
                    break
            content_items.append({"type": "bullet_list", "items": items})
            continue

        # Numbered list item
        num_match = re.match(r'^\d+\.\s+(.+)$', line)
        if num_match:
            items = [num_match.group(1)]
            i += 1
            while i < len(lines):
                next_num = re.match(r'^\d+\.\s+(.+)$', lines[i])
                if next_num:
                    items.append(next_num.group(1))
                    i += 1
                else:
                    break
            content_items.append({"type": "numbered_list", "items": items})
            continue

        # Blockquote
        quote_match = re.match(r'^>\s*(.+)$', line)
        if quote_match:
            quote_text = [quote_match.group(1)]
            i += 1
            while i < len(lines):
                next_quote = re.match(r'^>\s*(.*)$', lines[i])
                if next_quote:
                    quote_text.append(next_quote.group(1))
                    i += 1
                else:
                    break
            content_items.append({"type": "quote", "text": '\n'.join(quote_text)})
            continue

        # Horizontal rule
        if re.match(r'^[\-\*_]{3,}\s*$', line):
            content_items.append({"type": "horizontal_rule"})
            i += 1
            continue

        # Regular paragraph - collect consecutive non-empty lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            # Stop at empty line or special markdown syntax
            if not next_line.strip() or \
               re.match(r'^(#{1,6}|\*|\-|\+|\d+\.|\>|```|[\-\*_]{3,})', next_line):
                break
            para_lines.append(next_line)
            i += 1

        content_items.append({"type": "paragraph", "text": ' '.join(para_lines)})

    return content_items
