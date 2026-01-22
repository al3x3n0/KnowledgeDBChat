"""
Utility for parsing and manipulating docx templates.
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from loguru import logger
import re


class TemplateParser:
    """Utility for parsing and filling docx templates."""

    # Heading style patterns
    HEADING_STYLES = [
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
        'Title', 'Subtitle'
    ]

    @staticmethod
    def extract_sections(file_path: str) -> List[Dict[str, Any]]:
        """
        Extract sections from a docx template.

        Identifies sections by heading styles or patterns like "Section Name:" at paragraph start.

        Args:
            file_path: Path to the docx template file

        Returns:
            List of section dictionaries with:
            - title: Section heading text
            - level: Heading level (1-4, 0 for detected patterns)
            - placeholder_text: Existing content under this section
            - paragraph_indices: List of paragraph indices for this section
        """
        doc = Document(file_path)
        sections = []
        current_section: Optional[Dict[str, Any]] = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading
            is_heading = False
            heading_level = 0

            # Check by style
            if para.style and para.style.name:
                style_name = para.style.name
                if style_name in TemplateParser.HEADING_STYLES:
                    is_heading = True
                    if style_name == 'Title':
                        heading_level = 0
                    elif style_name == 'Subtitle':
                        heading_level = 1
                    elif style_name.startswith('Heading '):
                        try:
                            heading_level = int(style_name.replace('Heading ', ''))
                        except ValueError:
                            heading_level = 1

            # Check by pattern: "Something:" at the start of paragraph
            # But only if text is short (likely a section header)
            if not is_heading and len(text) < 100:
                # Pattern: starts with text followed by colon
                if re.match(r'^[A-Za-zА-Яа-яЁё0-9\s\-]+:\s*$', text):
                    is_heading = True
                    heading_level = 2

            # Check by bold formatting (common in templates)
            if not is_heading and para.runs:
                # If all runs are bold and text is short, likely a header
                all_bold = all(run.bold for run in para.runs if run.text.strip())
                if all_bold and len(text) < 100 and len(para.runs) > 0:
                    is_heading = True
                    heading_level = 2

            if is_heading:
                # Save previous section
                if current_section:
                    sections.append(current_section)

                # Start new section
                current_section = {
                    'title': text.rstrip(':').strip(),
                    'level': heading_level,
                    'placeholder_text': '',
                    'paragraph_indices': []
                }
            elif current_section:
                # Add content to current section
                current_section['placeholder_text'] += text + '\n'
                current_section['paragraph_indices'].append(i)

        # Don't forget the last section
        if current_section:
            sections.append(current_section)

        # Clean up placeholder text
        for section in sections:
            section['placeholder_text'] = section['placeholder_text'].strip()

        logger.info(f"Extracted {len(sections)} sections from template")
        return sections

    @staticmethod
    def fill_sections(
        template_path: str,
        sections_content: Dict[str, str],
        output_path: str
    ) -> str:
        """
        Fill template sections with generated content.

        Args:
            template_path: Path to the original template
            sections_content: Dictionary mapping section titles to generated content
            output_path: Path to save the filled document

        Returns:
            Path to the saved document
        """
        doc = Document(template_path)
        current_section_title: Optional[str] = None
        paragraphs_to_clear: List[int] = []
        first_content_para: Dict[str, int] = {}  # section_title -> first content paragraph index

        # First pass: identify sections and content paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            is_heading = TemplateParser._is_heading(para)

            if is_heading:
                section_title = text.rstrip(':').strip()
                if section_title in sections_content:
                    current_section_title = section_title
                    first_content_para[section_title] = None
                else:
                    current_section_title = None
            elif current_section_title:
                if first_content_para.get(current_section_title) is None:
                    first_content_para[current_section_title] = i
                else:
                    # Mark additional paragraphs for clearing
                    paragraphs_to_clear.append(i)

        # Second pass: fill content
        for section_title, content in sections_content.items():
            if section_title in first_content_para and first_content_para[section_title] is not None:
                idx = first_content_para[section_title]
                # Replace content of the first paragraph
                para = doc.paragraphs[idx]
                para.clear()
                # Add content, preserving paragraph breaks
                lines = content.split('\n')
                for j, line in enumerate(lines):
                    if j == 0:
                        run = para.add_run(line)
                    else:
                        # Add line break for subsequent lines within the same paragraph
                        para.add_run('\n' + line)

        # Clear extra paragraphs (mark as empty)
        for idx in paragraphs_to_clear:
            if idx < len(doc.paragraphs):
                doc.paragraphs[idx].clear()

        # Save the document
        doc.save(output_path)
        logger.info(f"Filled template saved to {output_path}")
        return output_path

    @staticmethod
    def _is_heading(para) -> bool:
        """Check if a paragraph is a heading."""
        text = para.text.strip()
        if not text:
            return False

        # Check by style
        if para.style and para.style.name:
            if para.style.name in TemplateParser.HEADING_STYLES:
                return True

        # Check by pattern
        if len(text) < 100 and re.match(r'^[A-Za-zА-Яа-яЁё0-9\s\-]+:\s*$', text):
            return True

        # Check by bold formatting
        if para.runs and len(text) < 100:
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            if all_bold and len(para.runs) > 0:
                return True

        return False

    @staticmethod
    def get_template_metadata(file_path: str) -> Dict[str, Any]:
        """
        Get metadata about the template document.

        Args:
            file_path: Path to the docx template

        Returns:
            Dictionary with template metadata
        """
        doc = Document(file_path)
        core_props = doc.core_properties

        return {
            'title': core_props.title or Path(file_path).stem,
            'author': core_props.author,
            'created': core_props.created,
            'modified': core_props.modified,
            'paragraph_count': len(doc.paragraphs),
            'word_count': sum(len(p.text.split()) for p in doc.paragraphs),
        }
