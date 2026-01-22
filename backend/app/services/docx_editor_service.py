"""
DOCX Editor Service - handles conversion between DOCX and HTML for web editing.
"""

import hashlib
import tempfile
import re
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path
from io import BytesIO
from loguru import logger

import mammoth
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup, NavigableString


class DocxEditorService:
    """Service for converting DOCX to HTML and back."""

    def __init__(self):
        # Custom style map for mammoth to preserve semantic structure
        self.style_map = """
            p[style-name='Heading 1'] => h1:fresh
            p[style-name='Heading 2'] => h2:fresh
            p[style-name='Heading 3'] => h3:fresh
            p[style-name='Heading 4'] => h4:fresh
            p[style-name='Heading 5'] => h5:fresh
            p[style-name='Heading 6'] => h6:fresh
            p[style-name='Title'] => h1.title:fresh
            p[style-name='Subtitle'] => h2.subtitle:fresh
            b => strong
            i => em
            u => u
            strike => s
        """

    async def docx_to_html(self, file_path: str) -> Dict[str, Any]:
        """
        Convert a DOCX file to HTML for editing.

        Args:
            file_path: Path to the DOCX file (local or temp file)

        Returns:
            Dict with html_content, warnings, and metadata
        """
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=self.style_map
                )

            html_content = result.value
            warnings = [msg.message for msg in result.messages]

            # Clean up the HTML for TipTap compatibility
            html_content = self._clean_html_for_editor(html_content)

            # Calculate content hash for version tracking
            content_hash = hashlib.sha256(html_content.encode()).hexdigest()[:16]

            # Extract document metadata
            doc = Document(file_path)
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "paragraph_count": len(doc.paragraphs),
            }

            return {
                "html_content": html_content,
                "warnings": warnings if warnings else None,
                "version": content_hash,
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"Failed to convert DOCX to HTML: {e}")
            raise ValueError(f"Failed to convert document: {str(e)}")

    def _clean_html_for_editor(self, html: str) -> str:
        """Clean and normalize HTML for TipTap editor."""
        soup = BeautifulSoup(html, "html.parser")

        # Ensure empty paragraphs have proper content for editing
        for p in soup.find_all("p"):
            if not p.get_text(strip=True):
                p.string = "\u00A0"  # Non-breaking space for empty paragraphs

        # Convert any remaining Word-specific elements
        # Remove empty spans
        for span in soup.find_all("span"):
            if not span.get_text(strip=True) and not span.find_all():
                span.decompose()

        return str(soup)

    async def html_to_docx(
        self,
        html_content: str,
        original_path: Optional[str] = None
    ) -> bytes:
        """
        Convert HTML back to DOCX format.

        Args:
            html_content: HTML string from the editor
            original_path: Optional path to original DOCX to preserve styles

        Returns:
            DOCX file as bytes
        """
        try:
            # Create a new document or use original as template
            if original_path and Path(original_path).exists():
                doc = Document(original_path)
                # Clear existing content but preserve styles
                for element in doc.element.body[:]:
                    doc.element.body.remove(element)
            else:
                doc = Document()
                self._setup_default_styles(doc)

            # Parse HTML
            soup = BeautifulSoup(html_content, "html.parser")

            # Convert HTML elements to DOCX
            self._convert_html_to_docx(soup, doc)

            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer.read()

        except Exception as e:
            logger.error(f"Failed to convert HTML to DOCX: {e}")
            raise ValueError(f"Failed to save document: {str(e)}")

    def _setup_default_styles(self, doc: Document) -> None:
        """Setup default styles for a new document."""
        # Styles are automatically available in python-docx
        pass

    def _convert_html_to_docx(self, soup: BeautifulSoup, doc: Document) -> None:
        """Convert parsed HTML to DOCX elements."""
        body = soup.body if soup.body else soup

        for element in body.children:
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    doc.add_paragraph(text)
            else:
                self._process_element(element, doc)

    def _process_element(self, element, doc: Document, parent_para=None) -> None:
        """Process a single HTML element and add to document."""
        tag_name = element.name.lower() if element.name else ""

        if tag_name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(tag_name[1])
            text = element.get_text()
            para = doc.add_paragraph(text, style=f"Heading {level}")

        elif tag_name == "p":
            para = doc.add_paragraph()
            self._process_inline_content(element, para)

        elif tag_name == "ul":
            self._process_list(element, doc, numbered=False)

        elif tag_name == "ol":
            self._process_list(element, doc, numbered=True)

        elif tag_name == "table":
            self._process_table(element, doc)

        elif tag_name == "blockquote":
            para = doc.add_paragraph(style="Quote")
            self._process_inline_content(element, para)

        elif tag_name == "br":
            if parent_para:
                parent_para.add_run("\n")

        elif tag_name in ["div", "section", "article"]:
            # Process children
            for child in element.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        doc.add_paragraph(text)
                else:
                    self._process_element(child, doc)

    def _process_inline_content(self, element, para) -> None:
        """Process inline content (text, bold, italic, etc.) within a paragraph."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip() or text == "\u00A0":
                    para.add_run(text)
            else:
                self._process_inline_element(child, para)

    def _process_inline_element(self, element, para) -> None:
        """Process inline formatting elements."""
        tag_name = element.name.lower() if element.name else ""
        text = element.get_text()

        if not text.strip() and text != "\u00A0":
            return

        run = para.add_run(text)

        # Apply formatting based on tag
        if tag_name in ["strong", "b"]:
            run.bold = True
        elif tag_name in ["em", "i"]:
            run.italic = True
        elif tag_name == "u":
            run.underline = True
        elif tag_name == "s" or tag_name == "strike":
            run.font.strike = True
        elif tag_name == "sub":
            run.font.subscript = True
        elif tag_name == "sup":
            run.font.superscript = True

        # Handle nested formatting
        if element.find(["strong", "b"]):
            run.bold = True
        if element.find(["em", "i"]):
            run.italic = True
        if element.find("u"):
            run.underline = True

    def _process_list(self, element, doc: Document, numbered: bool = False) -> None:
        """Process ordered or unordered list."""
        style = "List Number" if numbered else "List Bullet"

        for li in element.find_all("li", recursive=False):
            para = doc.add_paragraph(style=style)
            self._process_inline_content(li, para)

            # Handle nested lists
            nested_ul = li.find("ul", recursive=False)
            nested_ol = li.find("ol", recursive=False)
            if nested_ul:
                self._process_list(nested_ul, doc, numbered=False)
            if nested_ol:
                self._process_list(nested_ol, doc, numbered=True)

    def _process_table(self, element, doc: Document) -> None:
        """Process HTML table to DOCX table."""
        rows = element.find_all("tr")
        if not rows:
            return

        # Determine table dimensions
        num_rows = len(rows)
        num_cols = max(len(row.find_all(["td", "th"])) for row in rows)

        if num_cols == 0:
            return

        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            for col_idx, cell in enumerate(cells):
                if col_idx < num_cols:
                    table_cell = table.rows[row_idx].cells[col_idx]
                    # Clear default paragraph and add content
                    table_cell.text = ""
                    para = table_cell.paragraphs[0]
                    self._process_inline_content(cell, para)

                    # Bold for header cells
                    if cell.name == "th":
                        for run in para.runs:
                            run.bold = True

    def calculate_content_hash(self, content: str) -> str:
        """Calculate a hash for content versioning."""
        return hashlib.sha256(content.encode()).hexdigest()[:16]


# Global service instance
docx_editor_service = DocxEditorService()
